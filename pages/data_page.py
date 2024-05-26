import tkinter as tk
from tkinter import scrolledtext, messagebox
import os
import requests
from fpdf import FPDF
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from Helps.translations import translations
from Helps.utils import *
from datetime import datetime
from PIL import Image, ImageTk

SERVER_URL = "http://127.0.0.1:5000"

class DataPage(tk.Frame):
    def __init__(self, parent, app, user_id=None):
        super().__init__(parent)
        self.parent = parent  # Ссылка на родительский виджет
        self.app = app  # Ссылка на главное приложение
        self.user_id = user_id
        self.content = tk.Frame(self)  # Основная рамка для содержимого
        self.content.pack(expand=True, fill="both")
        self.buttons = {}  # Словарь для хранения кнопок
        self.current_language = 'ru'

        self.get_content()  # Инициализация контента страницы

    translation_key = 'data'

    def set_language(self, language):
        self.current_language = language
        # Обновление текста на всех созданных кнопках
        for class_id, button in self.buttons.items():
            button.config(text=translations[language]['class_data'].format(class_id))
        # Обновление текста других элементов
        if hasattr(self, 'data_button'):
            self.data_button.config(text=translations[language]['show_data'])
        if hasattr(self, 'graph1_button'):
            self.graph1_button.config(text=translations[language]['graph'] + " 1")
        if hasattr(self, 'graph2_button'):
            self.graph2_button.config(text=translations[language]['graph'] + " 2")
        if hasattr(self, 'graph3_button'):
            self.graph3_button.config(text=translations[language]['graph'] + " 3")
        if hasattr(self, 'pdfot4et'):
            self.pdfot4et.config(text=translations[language]['pdfot4et'])
        if hasattr(self, 'word_report_button'):
            self.word_report_button.config(text=translations[language]['word_report_button'])
        if hasattr(self, 'back_button'):
            self.back_button.config(text=translations[language]['back_to_menu'])
        self.update_idletasks()  # Обновление интерфейса

    def show_class_data(self, class_id):
        """Отображение данных класса или графиков."""
        self.content.destroy()
        self.content = tk.Frame(self)
        self.content.pack(expand=True, fill="both")

        self.data_text = scrolledtext.ScrolledText(self.content, wrap=tk.WORD, width=50, height=10)
        self.data_text.pack(pady=10)

        button_frame = tk.Frame(self.content)
        button_frame.pack(pady=10)

        # Создание кнопок с правильным текстом из словаря переводов
        self.data_button = tk.Button(button_frame, text=translations[self.current_language]['show_data'],
                                     command=lambda: self.fetch_and_display_data(class_id))
        self.data_button.pack(side="left", padx=5)

        self.graph1_button = tk.Button(button_frame, text=translations[self.current_language]['graph'] + " 1",
                                       command=lambda: self.display_graph(class_id, 1))
        self.graph1_button.pack(side="left", padx=5)

        self.graph2_button = tk.Button(button_frame, text=translations[self.current_language]['graph'] + " 2",
                                       command=lambda: self.display_graph(class_id, 2))
        self.graph2_button.pack(side="left", padx=5)

        self.graph3_button = tk.Button(button_frame, text=translations[self.current_language]['graph'] + " 3",
                                       command=lambda: self.display_graph(class_id, 3))
        self.graph3_button.pack(side="left", padx=5)

        self.pdfot4et = tk.Button(button_frame,text=translations[self.current_language]['pdfot4et'],
                                  command=lambda: self.save_report_to_pdf(class_id))
        self.pdfot4et.pack(side="left", padx=5)

        self.word_report_button = tk.Button(button_frame, text=translations[self.current_language]['word_report_button'],
                                            command=lambda: self.save_report_to_word(class_id))
        self.word_report_button.pack(side="left", padx=5)

        self.back_button = tk.Button(self.content, text=translations[self.current_language]['back_to_menu'],
                                     command=self.back_to_menu)
        self.back_button.pack(pady=10)

    def fetch_and_display_data(self, class_id):
        """Получение и отображение данных о последних 10 детекциях и данных последнего обработанного видео."""
        try:
            response = requests.get(f"{SERVER_URL}/get_class_data",
                                    params={"user_id": self.user_id, "class_id": class_id})
            if response.status_code == 200:
                rows = response.json()["data"]
                display_text = f"{translations[self.app.current_language]['last_10_detections'].format(class_id=class_id)}\n\n"
                for row in rows:
                    detection_date = datetime.strptime(row['detection_date'], "%a, %d %b %Y %H:%M:%S GMT")
                    formatted_date = detection_date.strftime("%Y-%m-%d %H:%M:%S")
                    display_text += f"{translations[self.app.current_language]['Datae']}: {formatted_date}, {translations[self.app.current_language]['Caounte']}: {row['people_count']}\n"
                self.data_text.delete(1.0, tk.END)
                self.data_text.insert(tk.END, display_text)
            else:
                messagebox.showerror(translations[self.app.current_language]['database_error'],
                                     translations[self.app.current_language]['fetch_data_error'].format(
                                         error=response.json().get("error")))
        except Exception as e:
            messagebox.showerror(translations[self.app.current_language]['database_error'],
                                 translations[self.app.current_language]['fetch_data_error'].format(error=e))

    def display_graph(self, class_id, graph_number):
        """Отображение выбранного графика."""
        try:
            response = requests.get(f"{SERVER_URL}/get_graph_path",
                                    params={"user_id": self.user_id, "class_id": class_id,
                                            "graph_number": graph_number})
            if response.status_code == 200:
                graph_path = response.json()["graph_path"]
                img = Image.open(graph_path)
                img = img.resize((500, 400), Image.LANCZOS)  # Адаптация размера изображения
                imgtk = ImageTk.PhotoImage(image=img)

                # Удаляем предыдущий виджет графика, если он существует
                if hasattr(self, 'graph_label'):
                    self.graph_label.destroy()

                # Создаем новый виджет графика
                self.graph_label = tk.Label(self.content, image=imgtk)
                self.graph_label.image = imgtk
                self.graph_label.pack(pady=10)
            else:
                messagebox.showerror(translations[self.app.current_language]['error'],
                                     translations[self.app.current_language]['graph_not_found_error'])
        except Exception as e:
            messagebox.showerror(translations[self.app.current_language]['error'],
                                 translations[self.app.current_language]['fetch_data_error'].format(error=e))


    def save_report_to_pdf(self, class_id):
        """Сохранение отчета в PDF файл с поддержкой Unicode и таблицей с бордюрами."""
        report_folder = 'reports/PDF format'
        if not os.path.exists(report_folder):
            os.makedirs(report_folder)

        username, position_name = self.get_user_details(self.user_id)
        current_date = datetime.now().strftime("%Y%m%d_%H%M%S")

        pdf_file_name = f"{translations[self.current_language]['report']}_{class_id}_{current_date}_{username}.pdf"
        pdf_file_path = os.path.join(report_folder, pdf_file_name)

        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('DejaVu', '', 'reports/shrift/DejaVuSans.ttf', uni=True)
        pdf.set_font('DejaVu', '', 12)

        pdf.cell(200, 10, txt=f"{translations[self.current_language]['username']}: {username}", ln=1)
        pdf.cell(200, 10, txt=f"{translations[self.current_language]['position']}: {position_name}", ln=1)
        pdf.cell(200, 10, txt=f"{translations[self.current_language]['ncam']}: {class_id}", ln=1)


        # Заголовок таблицы
        pdf.cell(30, 10, txt=translations[self.current_language]['nomerz'], border=1)
        pdf.cell(85, 10, txt=translations[self.current_language]['Date_of_discovery'], border=1)
        pdf.cell(85, 10, txt=translations[self.current_language]['Number_of_people'], border=1)
        pdf.ln(10)  # Переход на новую строку

        # Добавление данных
        data = self.fetch_data_for_ot4et(class_id)
        counter = 1
        for index, row in enumerate(data):
            try:
                detection_date = row.get('detection_date', '')
                people_count = row.get('people_count', '')

                if isinstance(detection_date, str):  # Проверка, является ли дата строкой
                    try:
                        detection_date = datetime.strptime(detection_date,
                                                           '%a, %d %b %Y %H:%M:%S %Z')  # Используем формат даты, соответствующий данным от сервера
                    except ValueError:
                        detection_date = datetime.now()  # В случае неудачи используем текущую дату

                formatted_date = detection_date.strftime('%Y-%m-%d %H:%M:%S')

                pdf.cell(30, 10, txt=str(counter), border=1)
                pdf.cell(85, 10, txt=formatted_date, border=1)
                pdf.cell(85, 10, txt=str(people_count), border=1)
                pdf.ln(10)
                counter += 1
            except Exception as e:
                messagebox.showerror(translations[self.current_language]['error'],
                                     translations[self.current_language]['fetch_data_error'].format(error=e))
                return

        # Добавление графиков
        for i in range(1, 4):
            graph_path = self.load_graph_path_from_db(class_id, i)
            if graph_path:
                pdf.add_page()
                pdf.image(graph_path, x=10, y=10, w=180)

        pdf.output(pdf_file_path)
        messagebox.showinfo(translations[self.current_language]['report_save'], f"{translations[self.current_language]['report_successfully_saved_to']}: {pdf_file_path}")

    def save_report_to_word(self, class_id):
        """Сохранение отчета в Word файл с графиками и форматированными датами."""
        report_folder = 'reports/Word format'
        if not os.path.exists(report_folder):
            os.makedirs(report_folder)

        username, position_name = self.get_user_details(self.user_id)
        current_date = datetime.now().strftime("%Y%m%d_%H%M%S")

        doc_file_name = f"{translations[self.current_language]['report']}_{class_id}_{current_date}_{username}.docx"
        doc_file_path = os.path.join(report_folder, doc_file_name)

        doc = Document()
        doc.add_heading(translations[self.current_language]['report'], level=1)

        doc.add_paragraph(f"{translations[self.current_language]['username']}: {username}")
        doc.add_paragraph(f"{translations[self.current_language]['position']}: {position_name}")
        doc.add_paragraph(f"{translations[self.current_language]['ncam']}: {class_id}")

        records = self.fetch_data_for_ot4et(class_id)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = translations[self.current_language]['nomerz']
        hdr_cells[1].text = translations[self.current_language]['Date_of_discovery']
        hdr_cells[2].text = translations[self.current_language]['Number_of_people']

        for index, record in enumerate(records):
            row_cells = table.add_row().cells
            detection_date = record.get('detection_date', '')
            people_count = record.get('people_count', '')

            if isinstance(detection_date, str):  # Проверка, является ли дата строкой
                try:
                    detection_date = datetime.strptime(detection_date,
                                                       '%a, %d %b %Y %H:%M:%S %Z')  # Используем формат даты, соответствующий данным от сервера
                except ValueError:
                    detection_date = datetime.now()  # В случае неудачи используем текущую дату

            formatted_date = detection_date.strftime("%Y-%m-%d %H:%M:%S")
            row_cells[0].text = str(index + 1)
            row_cells[1].text = formatted_date
            row_cells[2].text = str(people_count)
            for cell in row_cells:
                self.set_cell_border(cell, top={"sz": 12, "val": "single", "space": "0"},
                                     bottom={"sz": 12, "val": "single"})

        for i in range(1, 4):
            graph_path = self.load_graph_path_from_db(class_id, i)
            if graph_path:
                # Ensure each graph starts on a new page
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                doc.add_picture(graph_path, width=Pt(400))

        doc.save(doc_file_path)
        messagebox.showinfo(translations[self.current_language]['report_save'],
                            f"{translations[self.current_language]['report_successfully_saved_to']}: {doc_file_path}")

    def get_user_details(self, user_id):
        """Получение имени пользователя и перевода названия его должности по user_id."""
        try:
            response = requests.get(f"{SERVER_URL}/get_user_details", params={"user_id": user_id})
            if response.status_code == 200:
                user_info = response.json()
                username = user_info["username"]
                translated_position = translations[self.current_language]['positions'].get(user_info["position"],
                                                                                           user_info["position"])
                return username, translated_position
            else:
                return translations[self.current_language]['Unknown'], translations[self.current_language]['Unknown']
        except Exception as e:
            messagebox.showerror(translations[self.current_language]['database_error'],
                                 translations[self.current_language]['fetch_data_error'].format(error=e))
            return translations[self.current_language]['Unknown'], translations[self.current_language]['Unknown']

    def fetch_data_for_ot4et(self, class_id):
        """Запрос данных для отчета, начиная с последней обработанной видеосессии."""
        try:
            response = requests.get(f"{SERVER_URL}/get_class_data",
                                    params={"user_id": self.user_id, "class_id": class_id})
            if response.status_code == 200:
                data = response.json()["data"]
                # Преобразование строковой даты в объект datetime
                for row in data:
                    row['detection_date'] = datetime.strptime(row['detection_date'], "%a, %d %b %Y %H:%M:%S GMT")
                return data
            else:
                return []  # Если не удалось получить данные, возвращаем пустой список
        except Exception as e:
            messagebox.showerror(translations[self.current_language]['database_error'],
                                 translations[self.current_language]['fetch_data_error'].format(error=e))
            return []

    def set_cell_border(self, cell, **kwargs):
        """
        Set cell's border
        Usage:

        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existance, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def load_graph_path_from_db(self, class_id, graph_number):
        """Загрузка пути к графику из базы данных, соответствующего последнему обработанному видео."""
        try:
            response = requests.get(f"{SERVER_URL}/get_graph_path",
                                    params={"user_id": self.user_id, "class_id": class_id,
                                            "graph_number": graph_number})
            if response.status_code == 200:
                return response.json()["graph_path"]
            else:
                return None
        except Exception as e:
            messagebox.showerror(translations[self.current_language]['error'],
                                 translations[self.current_language]['fetch_data_error'].format(error=e))
            return None

    def create_button(self, index):
        """Создание кнопки для аудитории и сохранение ссылки на неё."""
        button = tk.Button(self.content, text=translations[self.current_language]['class_data'].format(index),
                           command=lambda i=index: self.show_class_data(i))
        self.buttons[index] = button  # Сохраняем ссылку на кнопку
        return button

    def get_user_position_id(self):
        """Получение ID должности пользователя из базы данных."""
        if self.user_id is None:
            print("user_id не установлен.")
            return None

        position_id = None
        try:
            conn = connect_db()
            if conn is None:
                print("Не удалось подключиться к базе данных.")
                return None

            cursor = conn.cursor()
            cursor.execute("SELECT position_id FROM users WHERE user_id = %s;", (self.user_id,))
            result = cursor.fetchone()
            if result:
                position_id = result[0]
            else:
                print(f"Не найден position_id для user_id {self.user_id}.")
            cursor.close()
        except Exception as e:
            print(f"Ошибка при получении ID должности: {e}")
        finally:
            if conn:
                conn.close()

        return position_id

    def generate_buttons(self):
        """Генерация кнопок для аудиторий на основе доступных данных по должности пользователя."""
        try:
            response = requests.get(f"{SERVER_URL}/get_class_ids", params={"user_id": self.user_id})
            if response.status_code == 200:
                class_ids = response.json()["class_ids"]
                for class_id in class_ids:
                    button = self.create_button(class_id)
                    button.pack(anchor="w", pady=(5, 0), padx=10, fill="x")
            else:
                messagebox.showerror(translations[self.current_language]['database_error'],
                                     translations[self.current_language]['fetch_data_error'].format(
                                         error=response.json().get("error")))
        except Exception as e:
            messagebox.showerror(translations[self.current_language]['database_error'],
                                 translations[self.current_language]['fetch_data_error'].format(error=e))

    def back_to_menu(self):
        """Возвращение в меню."""
        self.content.destroy()
        self.content = tk.Frame(self)
        self.content.pack(expand=True, fill="both")
        self.get_content()

    def get_content(self):
        """Получение контента страницы."""
        self.generate_buttons()
        return self.content
